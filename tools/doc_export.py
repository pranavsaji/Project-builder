# tools/doc_export.py
from __future__ import annotations

import io
import re
import zipfile
from typing import Any, Dict, List, Optional

# Optional docx dependency
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

LANG_LABEL = {
    "python": "python", "markdown": "md", "yaml": "yaml", "json": "json", "toml": "toml", "ini": "ini",
    "xml": "xml", "csv": "csv", "text": "", "html": "html", "css": "css", "scss": "scss", "sass": "sass",
    "javascript": "javascript", "typescript": "typescript", "tsx": "tsx", "jsx": "jsx", "bash": "bash",
    "sql": "sql", "protobuf": "protobuf", "docker": "docker",
}

def fence(lang: str) -> str:
    return LANG_LABEL.get((lang or "text").lower(), "")

def build_markdown_document(files: List[Dict], title: str, base_path: str) -> str:
    parts = [f"# {title}", "", f"_Base:_ `{base_path}`", ""]
    for f in files:
        lang = fence(f.get("language") or "text")
        rel = f.get("rel_path", "")
        parts.append(f"## `{rel}`")
        parts.append("")
        parts.append(f"```{lang}".rstrip())
        parts.append(f.get("content", ""))
        parts.append("```")
        parts.append("")  # spacer
    return "\n".join(parts)


# -----------------------------
# XML-safe DOCX export helpers
# -----------------------------
_XML_INVALID = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")  # allow \t \n \r

def _xml_safe_text(text: Any) -> str:
    """Return a string with characters invalid in XML 1.0 removed/replaced."""
    if text is None:
        return ""
    if isinstance(text, bytes):
        # decode binary safely; NULLs and other invalids will be scrubbed below
        text = text.decode("utf-8", "replace")
    else:
        text = str(text)
    return _XML_INVALID.sub(" ", text)


def build_docx_document(files: List[Dict], title: str) -> Optional[bytes]:
    """
    Build a .docx with headings per file and code-like body text.
    This is resilient to binary/garbled inputs by sanitizing text before writing.
    """
    if not DOCX_AVAILABLE:
        return None

    doc = Document()
    doc.add_heading(title, level=0)

    for f in files:
        rel = f.get("rel_path") or f.get("path") or f.get("name") or "File"
        doc.add_heading(str(rel), level=1)

        raw = f.get("content", "")
        safe = _xml_safe_text(raw)

        if not safe and raw:
            safe = "[content skipped: binary or contained unsupported characters]"

        p = doc.add_paragraph()
        run = p.add_run(safe)
        font = run.font
        font.name = "Courier New"
        font.size = Pt(10)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_zip_of_sources(files: List[Dict]) -> bytes:
    """
    Build a zip of the (selected) file contents. Accepts str or bytes content.
    """
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for f in files:
            rel = f.get("rel_path", "file.txt")
            content = f.get("content", "")
            zf.writestr(rel, content)
    return bio.getvalue()
